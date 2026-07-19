from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import re
import time
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


SCHEMA_VERSION = "5.8"
SKILL_VERSION = "5.8.3"
NUMBER_PATTERN = re.compile(
    r"(?<![\d.])[-+]?\d+(?:\.\d+)?(?:万亿元|千亿元|百亿元|亿元|万元|万人次|人次|万|亿|%|％|年|月|日)?"
)


def _load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def _write_json_atomic(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    temporary.replace(path)


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _numbers(values: list[tuple[str, str]]) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    for location, text in values:
        for match in NUMBER_PATTERN.finditer(text):
            value = match.group(0)
            if value:
                result.append({"value": value, "location": location, "context": text})
    return result


def parse_docx(path: str | Path) -> dict[str, Any]:
    source = Path(path).resolve(strict=True)
    document = Document(str(source))
    paragraphs: list[dict[str, Any]] = []
    searchable: list[tuple[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style is not None else ""
        record = {"index": index, "text": text, "style": style}
        paragraphs.append(record)
        searchable.append((f"paragraph:{index}", text))
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        rows: list[list[str]] = []
        for row_index, row in enumerate(table.rows):
            values = [cell.text.strip() for cell in row.cells]
            rows.append(values)
            for column_index, value in enumerate(values):
                if value:
                    searchable.append((f"table:{table_index}/r{row_index}/c{column_index}", value))
        tables.append({"index": table_index, "rows": rows})
    media: list[dict[str, Any]] = []
    relationships: list[dict[str, str]] = []
    with zipfile.ZipFile(source) as archive:
        for name in sorted(archive.namelist()):
            if name.startswith("word/media/") and not name.endswith("/"):
                payload = archive.read(name)
                media.append(
                    {
                        "path": name,
                        "name": Path(name).name,
                        "size": len(payload),
                        "sha256": hashlib.sha256(payload).hexdigest(),
                    }
                )
        rel_name = "word/_rels/document.xml.rels"
        if rel_name in archive.namelist():
            relation_text = archive.read(rel_name).decode("utf-8", errors="replace")
            for match in re.finditer(
                r'<Relationship\b[^>]*\bId="([^"]+)"[^>]*\bType="([^"]+)"[^>]*\bTarget="([^"]+)"',
                relation_text,
            ):
                relationships.append({"id": match.group(1), "type": match.group(2), "target": match.group(3)})
    headings = [
        {"text": item["text"], "style": item["style"], "paragraph_index": item["index"]}
        for item in paragraphs
        if str(item["style"]).lower().startswith("heading") or str(item["style"]).startswith("标题")
    ]
    return {
        "path": str(source),
        "name": source.name,
        "parser": "structural_docx",
        "paragraphs": paragraphs,
        "headings": headings,
        "tables": tables,
        "media": media,
        "relationships": relationships,
        "evidence_index": {
            "numbers": _numbers(searchable),
            "table_count": len(tables),
            "media_count": len(media),
        },
    }


def parse_text(path: str | Path) -> dict[str, Any]:
    source = Path(path).resolve(strict=True)
    text = source.read_text(encoding="utf-8-sig")
    lines = [{"index": index, "text": line} for index, line in enumerate(text.splitlines()) if line.strip()]
    return {
        "path": str(source),
        "name": source.name,
        "parser": "utf8_text",
        "paragraphs": lines,
        "headings": [],
        "tables": [],
        "media": [],
        "relationships": [],
        "evidence_index": {
            "numbers": _numbers([(f"line:{item['index']}", item["text"]) for item in lines]),
            "table_count": 0,
            "media_count": 0,
        },
    }


def parse_source(path: str | Path) -> dict[str, Any]:
    source = Path(path).expanduser().resolve(strict=True)
    if not source.is_file():
        raise FileNotFoundError(source)
    suffix = source.suffix.lower()
    if suffix == ".docx":
        return parse_docx(source)
    if suffix in {".txt", ".md", ".csv", ".json"}:
        return parse_text(source)
    return {
        "path": str(source),
        "name": source.name,
        "parser": "specialized_parser_required",
        "paragraphs": [],
        "headings": [],
        "tables": [],
        "media": [],
        "relationships": [],
        "evidence_index": {"numbers": [], "table_count": 0, "media_count": 0},
        "requires_specialized_parser": True,
    }


def _extract_docx_media(project: Path, parsed: dict[str, Any], source_index: int) -> None:
    if parsed.get("parser") != "structural_docx" or not parsed.get("media"):
        return
    source = Path(str(parsed["path"]))
    destination_dir = project / ".build" / "source_media" / f"{source_index:02d}_{source.stem}"
    destination_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(source) as archive:
        for item in parsed["media"]:
            archive_path = str(item["path"])
            destination = destination_dir / Path(archive_path).name
            temporary = destination.with_suffix(destination.suffix + ".tmp")
            temporary.write_bytes(archive.read(archive_path))
            temporary.replace(destination)
            item["extracted_path"] = destination.relative_to(project).as_posix()


def _brief_sources(project_dir: Path) -> list[Path]:
    brief = json.loads((project_dir / "project_brief.json").read_text(encoding="utf-8"))
    raw_sources = brief.get("source_files")
    if not isinstance(raw_sources, list) or not raw_sources:
        raise ValueError("V5.8.3 project_brief.json requires a non-empty source_files list")
    sources: list[Path] = []
    for raw_path in raw_sources:
        if not isinstance(raw_path, str) or not raw_path.strip():
            raise ValueError("every source_files entry must be a non-empty UTF-8 path string")
        source = Path(raw_path).expanduser()
        try:
            resolved = source.resolve(strict=True)
        except FileNotFoundError as exc:
            raise FileNotFoundError(source) from exc
        if not resolved.is_file():
            raise FileNotFoundError(resolved)
        sources.append(resolved)
    return sources


def ingest_project_sources(project_dir: str | Path) -> dict[str, Any]:
    project = Path(project_dir).resolve()
    brief = json.loads((project / "project_brief.json").read_text(encoding="utf-8"))
    schema_version = str(brief.get("schema_version", SCHEMA_VERSION))
    skill_version = "5.9.0" if schema_version == "5.9" else SKILL_VERSION
    start = time.time()
    sources = _brief_sources(project)
    located = time.time()
    source_cache = _load_module(
        "standard_report_v583_source_cache",
        Path(__file__).with_name("v58_source_cache.py"),
    )
    extract_path = project / ".build" / "source_extract.json"
    cached_digest = source_cache.load_source_digest(project, sources)
    if cached_digest is not None and extract_path.is_file():
        try:
            extract = json.loads(extract_path.read_text(encoding="utf-8"))
        except (OSError, ValueError, json.JSONDecodeError):
            extract = None
        if isinstance(extract, dict) and extract.get("source_set_sha256") == cached_digest.get("source_set_sha256"):
            end = time.time()
            report = {
                "schema_version": schema_version,
                "skill_version": skill_version,
                "ok": True,
                "cache_hit": True,
                "duration_seconds": round(end - start, 3),
                "stages": [
                    {
                        "stage": "source_locate",
                        "start_epoch": start,
                        "end_epoch": located,
                        "cache_hit": False,
                    },
                    {
                        "stage": "source_parse",
                        "start_epoch": located,
                        "end_epoch": end,
                        "cache_hit": True,
                    },
                    {
                        "stage": "evidence_index",
                        "start_epoch": end,
                        "end_epoch": end,
                        "cache_hit": True,
                    },
                ],
                "sources": [
                    {
                        "path": item["path"],
                        "name": item["name"],
                        "parser": parsed.get("parser", "cached"),
                    }
                    for item, parsed in zip(cached_digest["sources"], extract.get("sources", []))
                ],
            }
            _write_json_atomic(project / ".build" / "source_ingest_report.json", report)
            return report
    parsed_sources = [parse_source(source) for source in sources]
    for source_index, parsed_source in enumerate(parsed_sources, start=1):
        _extract_docx_media(project, parsed_source, source_index)
    parsed = time.time()
    source_set_hash = source_cache.source_set_sha256(sources)
    extract = {
        "schema_version": schema_version,
        "skill_version": skill_version,
        "source_set_sha256": source_set_hash,
        "sources": parsed_sources,
    }
    _write_json_atomic(extract_path, extract)
    source_cache.write_source_digest(project, sources, extract, schema_version=schema_version)
    end = time.time()
    report = {
        "schema_version": schema_version,
        "skill_version": skill_version,
        "ok": True,
        "cache_hit": False,
        "duration_seconds": round(end - start, 3),
        "stages": [
            {
                "stage": "source_locate",
                "start_epoch": start,
                "end_epoch": located,
                "cache_hit": False,
            },
            {
                "stage": "source_parse",
                "start_epoch": located,
                "end_epoch": parsed,
                "cache_hit": False,
            },
            {
                "stage": "evidence_index",
                "start_epoch": parsed,
                "end_epoch": end,
                "cache_hit": False,
            },
        ],
        "sources": [
            {"path": item["path"], "name": item["name"], "parser": item["parser"]}
            for item in parsed_sources
        ],
    }
    _write_json_atomic(project / ".build" / "source_ingest_report.json", report)
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Parse and cache V5.8.3 source files without shell path rewriting.")
    parser.add_argument("project", type=Path)
    args = parser.parse_args()
    print(json.dumps(ingest_project_sources(args.project), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
