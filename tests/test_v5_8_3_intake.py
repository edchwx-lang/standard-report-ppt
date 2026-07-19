from __future__ import annotations

import hashlib
import importlib.util
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image


SKILL = Path(__file__).resolve().parents[1]


def load_module(name: str, path: Path):
    if not path.is_file():
        raise AssertionError(f"required V5.8.3 module is missing: {path.name}")
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_docx(path: Path) -> None:
    image = path.with_suffix(".png")
    Image.new("RGB", (80, 48), "#3F628F").save(image)
    document = Document()
    document.add_heading("中国农文旅市场", level=1)
    document.add_paragraph("2025年市场规模达到6.30万亿元，同比增长8.2%。")
    table = document.add_table(rows=3, cols=3)
    for column, value in enumerate(("年份", "市场规模", "增速")):
        table.cell(0, column).text = value
    for row, values in enumerate((("2024", "5.82", "12.1%"), ("2025", "6.30", "8.2%")), start=1):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    document.add_picture(str(image))
    document.save(path)


class V583SourceIngestTests(unittest.TestCase):
    def test_chinese_docx_path_is_structurally_parsed_and_cached(self):
        ingest = load_module(
            "v583_source_ingest",
            SKILL / "scripts" / "v583_source_ingest.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            project = root / "项目 空格（测试）"
            project.mkdir()
            source = root / "P1 农文旅（正式）.docx"
            make_docx(source)
            (project / "project_brief.json").write_text(
                json.dumps(
                    {
                        "schema_version": "5.8",
                        "pipeline_revision": "5.8.3",
                        "requested_page_count": 1,
                        "production_mode": "blueprint",
                        "blueprint_engine": "direct",
                        "source_files": [str(source)],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            first = ingest.ingest_project_sources(project)
            second = ingest.ingest_project_sources(project)
            extract = json.loads(
                (project / ".build" / "source_extract.json").read_text(encoding="utf-8")
            )
            extracted_media = (
                project
                / extract["sources"][0]["media"][0]["extracted_path"]
            ).is_file()
        self.assertFalse(first["cache_hit"])
        self.assertTrue(second["cache_hit"])
        self.assertEqual("structural_docx", first["sources"][0]["parser"])
        payload = extract["sources"][0]
        self.assertEqual("中国农文旅市场", payload["paragraphs"][0]["text"])
        self.assertEqual("6.30", payload["tables"][0]["rows"][2][1])
        self.assertEqual(1, len(payload["media"]))
        self.assertTrue(extracted_media)
        self.assertGreaterEqual(len(payload["relationships"]), 1)
        self.assertTrue(any(item["value"] == "6.30万亿元" for item in payload["evidence_index"]["numbers"]))

    def test_missing_exact_source_does_not_silently_select_another_file(self):
        ingest = load_module(
            "v583_source_ingest_missing",
            SKILL / "scripts" / "v583_source_ingest.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            project = root / "project"
            project.mkdir()
            (root / "同名.docx").write_bytes(b"not the requested file")
            (project / "project_brief.json").write_text(
                json.dumps(
                    {
                        "schema_version": "5.8",
                        "pipeline_revision": "5.8.3",
                        "requested_page_count": 1,
                        "production_mode": "fast",
                        "source_files": [str(root / "missing" / "同名.docx")],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            with self.assertRaises(FileNotFoundError):
                ingest.ingest_project_sources(project)


class V583AuthoringTests(unittest.TestCase):
    def test_one_authoring_bundle_materializes_all_manifests_and_refreshes_benchmark(self):
        authoring = load_module(
            "v583_authoring",
            SKILL / "scripts" / "v583_authoring.py",
        )
        benchmark = load_module(
            "v583_benchmark",
            SKILL / "scripts" / "v58_text_benchmark.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            project = Path(directory)
            build = project / ".build"
            drafts = build / "design_drafts"
            drafts.mkdir(parents=True)
            draft = drafts / "S01.png"
            Image.new("RGB", (1600, 900), "white").save(draft)
            bundle = {
                "schema_version": "5.8",
                "slides": [
                    {
                        "slide_id": "S01",
                        "chapter": "一、行业概览",
                        "title": "标题A",
                        "core_points": ["核心判断"],
                        "source": "资料来源：测试",
                        "visual_route": {"data_kind": "qualitative", "qualitative_form": "narrative"},
                        "modules": [{"module_id": "M01"}],
                        "primary_visual_module_id": "M01",
                        "evidence_inventory": [
                            {
                                "evidence_id": "E01",
                                "statement": "原始证据",
                                "priority": "must_keep",
                                "module_id": "M01",
                            }
                        ],
                    }
                ],
                "page_specs": {
                    "S01": {
                        "elements": [
                            {
                                "type": "text_card",
                                "role": "primary_evidence",
                                "box": [0.2, 0.2, 11.8, 3.8],
                                "title": "标题A",
                                "body": "原始证据",
                            }
                        ]
                    }
                },
                "visual_manifest": {
                    "schema_version": "5.8",
                    "pages": {
                        "S01": {
                            "design_draft_path": ".build/design_drafts/S01.png",
                            "imagegen_attempt_count": 1,
                            "transport_attempt_count": 1,
                        }
                    },
                },
            }
            (build / "authoring_bundle.json").write_text(
                json.dumps(bundle, ensure_ascii=False),
                encoding="utf-8",
            )
            authoring.materialize_project(project)
            first_benchmark = json.loads(
                (build / "blueprint_text_benchmark.json").read_text(encoding="utf-8")
            )
            bundle["slides"][0]["title"] = "标题B"
            (build / "authoring_bundle.json").write_text(
                json.dumps(bundle, ensure_ascii=False),
                encoding="utf-8",
            )
            authoring.materialize_project(project)
            second_benchmark = json.loads(
                (build / "blueprint_text_benchmark.json").read_text(encoding="utf-8")
            )
            manifest = json.loads((build / "visual_manifest.json").read_text(encoding="utf-8"))
            slides = json.loads((build / "slides.json").read_text(encoding="utf-8"))
            draft_digest = hashlib.sha256(draft.read_bytes()).hexdigest()
        self.assertEqual(draft_digest, manifest["pages"]["S01"]["design_draft_sha256"])
        self.assertEqual("标题B", slides[0]["title"])
        self.assertNotEqual(
            first_benchmark["pages"]["S01"]["canonical_sha256"],
            second_benchmark["pages"]["S01"]["canonical_sha256"],
        )
        self.assertEqual(
            benchmark.canonical_sha256(slides[0], bundle["page_specs"]["S01"]),
            second_benchmark["pages"]["S01"]["canonical_sha256"],
        )

    def test_automatic_materialization_rebinds_a_new_design_draft(self):
        pipeline = load_module(
            "v583_authoring_pipeline_rebind",
            SKILL / "scripts" / "project_pipeline.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            project = Path(directory)
            build = project / ".build"
            build.mkdir()
            (project / "project_brief.json").write_text(
                json.dumps(
                    {
                        "schema_version": "5.8",
                        "pipeline_revision": "5.8.3",
                        "requested_page_count": 1,
                        "production_mode": "blueprint",
                        "blueprint_engine": "direct",
                        "source_files": ["unused-in-this-unit-test"],
                    }
                ),
                encoding="utf-8",
            )
            bundle = {
                "schema_version": "5.8",
                "slides": [
                    {
                        "slide_id": "S01",
                        "chapter": "章节",
                        "title": "标题",
                        "core_points": ["判断"],
                        "source": "来源",
                    }
                ],
                "page_specs": {
                    "S01": {
                        "elements": [
                            {
                                "type": "text_card",
                                "box": [0.2, 0.2, 11.8, 3.8],
                                "title": "标题",
                                "body": "正文",
                            }
                        ]
                    }
                },
                "visual_manifest": {
                    "schema_version": "5.8",
                    "pages": {
                        "S01": {
                            "design_draft_path": ".build/design_drafts/S01.png",
                            "imagegen_attempt_count": 1,
                        }
                    },
                },
            }
            (build / "authoring_bundle.json").write_text(
                json.dumps(bundle, ensure_ascii=False),
                encoding="utf-8",
            )
            pipeline.materialize_project(project)
            draft = build / "design_drafts" / "S01.png"
            draft.parent.mkdir()
            Image.new("RGB", (1600, 900), "white").save(draft)
            pipeline._materialize_v583_if_present(project)
            manifest = json.loads(
                (build / "visual_manifest.json").read_text(encoding="utf-8")
            )
            draft_digest = hashlib.sha256(draft.read_bytes()).hexdigest()
        self.assertEqual(
            draft_digest,
            manifest["pages"]["S01"]["design_draft_sha256"],
        )


class V583LayoutAndTimingTests(unittest.TestCase):
    def test_layout_precheck_reports_unintended_asset_text_overlap_but_not_allowed_overlap(self):
        prebuild = load_module(
            "v583_layout_precheck",
            SKILL / "scripts" / "v58_prebuild.py",
        )
        specs = {
            "S01": {
                "elements": [
                    {"type": "asset", "asset_id": "A01", "box": [0.2, 0.2, 2.0, 1.5]},
                    {"type": "text", "text": "说明文字", "box": [1.0, 0.5, 3.0, 1.0]},
                    {"type": "rect", "box": [5.0, 0.2, 2.0, 1.0], "allow_overlap": True},
                    {"type": "text", "text": "允许叠放", "box": [5.2, 0.4, 1.0, 0.4], "allow_overlap": True},
                ]
            }
        }
        report = prebuild.diagnose_layout(specs)
        self.assertEqual("pass_with_warnings", report["status"])
        self.assertEqual(1, len(report["warnings"]))
        self.assertEqual("LAYOUT_OVERLAP", report["warnings"][0]["code"])

    def test_timing_keeps_front_stages_and_summarizes_wall_clock(self):
        timing = load_module(
            "v583_timing",
            SKILL / "scripts" / "v583_timing.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            project = Path(directory)
            timing.initialize_timing(project, started_epoch=100.0)
            timing.record_stage(
                project,
                "source_parse",
                105.0,
                115.0,
                ok=True,
                cache_hit=False,
                attempt_count=1,
            )
            timing.initialize_timing(project, started_epoch=120.0)
            timing.record_stage(
                project,
                "content_plan",
                116.0,
                140.0,
                ok=True,
            )
            payload = timing.summarize_timing(project, ended_epoch=150.0)
        self.assertEqual(["source_parse", "content_plan"], [item["stage"] for item in payload["stages"]])
        self.assertEqual(50.0, payload["wall_clock_seconds"])
        self.assertEqual(34.0, payload["active_seconds"])
        self.assertFalse(payload["stages"][0]["cache_hit"])

    def test_red_data_emphasis_role_alias_is_accepted(self):
        policy = load_module(
            "v583_visual_policy",
            SKILL / "scripts" / "v58_visual_policy.py",
        )
        spec = {
            "elements": [
                {
                    "type": "text",
                    "role": "data_emphasis",
                    "text": "6.30万亿元",
                    "color": "#C00000",
                    "box": [0.2, 0.2, 2.0, 0.5],
                },
                {"type": "rect", "fill": "#EDEDED", "box": [0.0, 1.0, 12.2, 0.5]},
            ]
        }
        self.assertEqual([], policy.validate_palette_and_visuals(spec, {}))

    def test_v583_init_ingests_sources_without_resetting_existing_timing(self):
        pipeline = load_module(
            "v583_pipeline_init",
            SKILL / "scripts" / "project_pipeline.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "材料 中文.docx"
            make_docx(source)
            project = root / "project"
            project.mkdir()
            (project / "project_brief.json").write_text(
                json.dumps(
                    {
                        "schema_version": "5.8",
                        "pipeline_revision": "5.8.3",
                        "requested_page_count": 1,
                        "production_mode": "fast",
                        "source_files": [str(source)],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            timing = load_module(
                "v583_pipeline_timing_for_init",
                SKILL / "scripts" / "v583_timing.py",
            )
            timing.initialize_timing(project, started_epoch=100.0)
            timing.record_stage(project, "skill_load", 100.0, 101.0, ok=True)
            result = pipeline.init_project(project)
            extract_exists = (project / ".build" / "source_extract.json").is_file()
            payload = json.loads(
                (project / ".build" / "pipeline_timing.json").read_text(encoding="utf-8")
            )
        self.assertEqual("5.8.3", result["skill_version"])
        self.assertTrue(extract_exists)
        self.assertEqual("skill_load", payload["stages"][0]["stage"])
        self.assertTrue(any(item["stage"] == "source_parse" for item in payload["stages"]))

    def test_v583_prebuild_writes_layout_precheck_report(self):
        pipeline = load_module(
            "v583_pipeline_prebuild",
            SKILL / "scripts" / "project_pipeline.py",
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            build = root / ".build"
            build.mkdir()
            (root / "project_brief.json").write_text(
                json.dumps(
                    {
                        "schema_version": "5.8",
                        "pipeline_revision": "5.8.3",
                        "requested_page_count": 1,
                        "production_mode": "fast",
                    }
                ),
                encoding="utf-8",
            )
            slides = [
                {
                    "slide_id": "S01",
                    "chapter": "章节",
                    "title": "标题",
                    "core_points": ["核心判断"],
                    "source": "来源",
                    "visual_route": {"data_kind": "qualitative", "qualitative_form": "narrative"},
                    "modules": [{"module_id": "M01"}],
                    "primary_visual_module_id": "M01",
                    "evidence_inventory": [
                        {
                            "evidence_id": "E01",
                            "statement": "证据",
                            "priority": "must_keep",
                            "module_id": "M01",
                        }
                    ],
                }
            ]
            page_specs = {
                "S01": {
                    "elements": [
                        {"type": "asset", "asset_id": "A01", "box": [0.2, 0.2, 2.0, 1.5]},
                        {
                            "type": "text",
                            "role": "primary_evidence",
                            "text": "重叠文字",
                            "box": [1.0, 0.5, 3.0, 1.0],
                        },
                    ]
                }
            }
            (build / "slides.json").write_text(json.dumps(slides, ensure_ascii=False), encoding="utf-8")
            (build / "page_specs.json").write_text(
                json.dumps(page_specs, ensure_ascii=False),
                encoding="utf-8",
            )
            result = pipeline.prebuild_project(root)
            report = json.loads(
                (build / "layout_precheck.json").read_text(encoding="utf-8")
            )
        self.assertTrue(result["ok"])
        self.assertTrue(any(item["code"] == "LAYOUT_OVERLAP" for item in report["warnings"]))
        self.assertEqual("5.8.3", result["skill_version"])


class V583DocumentationTests(unittest.TestCase):
    def test_skill_declares_v583_intake_fast_path_without_visual_contract_changes(self):
        skill = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        schema = (SKILL / "references" / "slide_spec_schema.md").read_text(encoding="utf-8")
        combined = skill + "\n" + schema
        self.assertIn("V5.8.3", combined)
        self.assertIn("source_files", combined)
        self.assertIn("authoring_bundle.json", combined)
        self.assertIn("layout_precheck.json", combined)
        self.assertIn("one successful ImageGen", skill)
        self.assertIn("Microsoft YaHei", skill)

    def test_v584_keeps_v583_imagegen_resources_and_versions_post_blueprint_resources(self):
        unchanged_paths = [
            SKILL / "prompts" / "imagegen_blueprint_prompt.md",
            SKILL / "references" / "company_visual_system.md",
            SKILL / "references" / "layout_and_chart_rules.md",
        ]
        for path in unchanged_paths:
            self.assertIn("V5.8.3", path.read_text(encoding="utf-8"), path.name)
        aligned_paths = [
            SKILL / "prompts" / "python_reconstruction_prompt.md",
            SKILL / "prompts" / "blueprint_alignment_prompt.md",
            SKILL / "references" / "ppt_quality_check_rules.md",
            SKILL / "references" / "python_reconstruction_rules.md",
        ]
        for path in aligned_paths:
            self.assertIn("V5.8.4", path.read_text(encoding="utf-8"), path.name)


if __name__ == "__main__":
    unittest.main()
