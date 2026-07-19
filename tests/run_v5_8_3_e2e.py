from __future__ import annotations

import hashlib
import importlib.util
import json
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw


SKILL = Path(__file__).resolve().parents[1]


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def main() -> None:
    pipeline = load_module("v583_e2e_pipeline", SKILL / "scripts" / "project_pipeline.py")
    pack = load_module("v583_e2e_pack", SKILL / "scripts" / "pack_delivery.py")

    with tempfile.TemporaryDirectory(prefix="标准报告_V583_") as directory:
        root = Path(directory)
        source = root / "P1 农文旅（端到端）.docx"
        document = Document()
        document.add_heading("中国文旅产业概况", level=1)
        document.add_paragraph("2025年居民出游总花费达到6.30万亿元，同比增长8.2%。")
        table = document.add_table(rows=4, cols=2)
        for row, values in enumerate(
            (("年份", "总花费"), ("2023", "4.91"), ("2024", "5.82"), ("2025", "6.30"))
        ):
            table.cell(row, 0).text = values[0]
            table.cell(row, 1).text = values[1]
        document.save(source)

        project = root / "V5.8.3 中文项目"
        project.mkdir()
        brief = {
            "schema_version": "5.8",
            "pipeline_revision": "5.8.3",
            "requested_page_count": 1,
            "production_mode": "blueprint",
            "blueprint_engine": "direct",
            "confirmation_source": "test_explicit",
            "source_files": [str(source)],
        }
        (project / "project_brief.json").write_text(
            json.dumps(brief, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        init_result = pipeline.init_project(project)

        draft = project / ".build" / "design_drafts" / "S01.png"
        image = Image.new("RGB", (1600, 900), "white")
        draw = ImageDraw.Draw(image)
        draw.rectangle((55, 70, 1545, 155), fill="#1E386B")
        draw.rectangle((80, 280, 1020, 735), fill="#F5F5F5", outline="#3F628F", width=5)
        draw.rectangle((1050, 280, 1515, 735), fill="#EDEDED", outline="#7391B3", width=5)
        for index, height in enumerate((120, 205, 285)):
            x = 180 + index * 250
            draw.rectangle((x, 670 - height, x + 105, 670), fill="#3F628F")
        draw.line((230, 555, 480, 475, 730, 365), fill="#C00000", width=8)
        image.save(draft)

        slides = [
            {
                "slide_id": "S01",
                "chapter": "一、中国文旅产业概况",
                "title": "2025年文旅消费修复至6.30万亿元",
                "core_points": ["■ 市场完成修复后进入常态增长阶段，服务消费升级继续支撑结构性增量。"],
                "source": "资料来源：V5.8.3 中文DOCX端到端测试",
                "page_number": 1,
                "density_profile": "adaptive",
                "visual_route": {
                    "data_kind": "time_series",
                    "data": [
                        {"label": "2023", "value": 4.91},
                        {"label": "2024", "value": 5.82},
                        {"label": "2025", "value": 6.30},
                    ],
                },
                "visual_brief": {
                    "primary_expression": "line_chart",
                    "visual_story": "Show the recovery path and the structural driver",
                    "supporting_visuals": [],
                },
                "modules": [{"module_id": "M01"}, {"module_id": "M02"}],
                "primary_visual_module_id": "M01",
                "evidence_inventory": [
                    {
                        "evidence_id": "E01",
                        "statement": "2025年居民出游总花费达到6.30万亿元",
                        "priority": "must_keep",
                        "module_id": "M01",
                    }
                ],
            }
        ]
        page_specs = {
            "S01": {
                "elements": [
                    {
                        "type": "line_chart",
                        "role": "primary_evidence",
                        "box": [0.2, 0.25, 7.7, 3.65],
                        "data": [
                            {"label": "2023", "value": 4.91},
                            {"label": "2024", "value": 5.82},
                            {
                                "label": "2025",
                                "value": 6.30,
                                "highlight": True,
                                "fill": "#C00000",
                            },
                        ],
                    },
                    {
                        "type": "text_card",
                        "box": [8.15, 0.25, 3.85, 3.65],
                        "title": "结构性驱动",
                        "body": "服务消费升级、下沉市场扩容与入境游恢复共同支撑增长。",
                        "title_fill": "#7391B3",
                        "body_fill": "#EDEDED",
                    },
                ]
            }
        }
        authoring_bundle = {
            "schema_version": "5.8",
            "slides": slides,
            "page_specs": page_specs,
            "visual_manifest": {
                "schema_version": "5.8",
                "pages": {
                    "S01": {
                        "design_draft_path": ".build/design_drafts/S01.png",
                        "imagegen_attempt_count": 1,
                        "transport_attempt_count": 1,
                        "visual_plan": [],
                        "visual_reviewed": True,
                        "observed_candidate_count": 0,
                        "candidate_count": 0,
                        "visuals": [],
                    }
                },
            },
        }
        (project / ".build" / "authoring_bundle.json").write_text(
            json.dumps(authoring_bundle, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        materialize_result = pipeline.materialize_project(project)
        output = project / "output" / "V583_农文旅_1页.pptx"
        result = pipeline.run_project(project, output)
        delivery = project / "delivery" / "V583_农文旅_1页.zip"
        pack.package_direct_delivery(
            project_dir=project,
            pptx_path=output,
            generator_path=project / "generate_deck.py",
            output_zip=delivery,
            desktop_dir=project / "unused_desktop",
        )

        quality = json.loads(
            (project / ".build" / "quality_report.json").read_text(encoding="utf-8")
        )
        asset_report = json.loads(
            (project / ".build" / "direct_asset_report.json").read_text(encoding="utf-8")
        )
        timing = json.loads(
            (project / ".build" / "pipeline_timing.json").read_text(encoding="utf-8")
        )
        stages = [item["stage"] for item in timing["stages"]]
        assert init_result["skill_version"] == "5.8.3"
        assert materialize_result["skill_version"] == "5.8.3"
        assert result["skill_version"] == "5.8.3"
        assert result["ok"] is True and result["blocker_count"] == 0
        assert quality["skill_version"] == "5.8.3"
        assert asset_report["skill_version"] == "5.8.3"
        assert (project / ".build" / "layout_precheck.json").is_file()
        for required in (
            "source_locate",
            "source_parse",
            "evidence_index",
            "canonical_materialize",
            "build",
            "render",
            "package",
        ):
            assert required in stages, (required, stages)
        assert timing["wall_clock_seconds"] >= timing["active_seconds"]
        assert (project / "blueprints" / "S01.png").read_bytes() == draft.read_bytes()
        assert output.is_file() and output.stat().st_size > 0
        assert delivery.is_file() and delivery.stat().st_size > 0
        with zipfile.ZipFile(delivery) as archive:
            assert set(archive.namelist()) == {
                output.name,
                "blueprints.zip",
                "py.zip",
            }
        print(
            json.dumps(
                {
                    "ok": True,
                    "skill_version": result["skill_version"],
                    "quality_status": result["quality_status"],
                    "warning_count": result["warning_count"],
                    "blocker_count": result["blocker_count"],
                    "source_parse_seconds": next(
                        item["duration_seconds"]
                        for item in timing["stages"]
                        if item["stage"] == "source_parse"
                    ),
                    "wall_clock_seconds": timing["wall_clock_seconds"],
                    "pptx_sha256": hashlib.sha256(output.read_bytes()).hexdigest(),
                    "pptx_bytes": output.stat().st_size,
                    "zip_bytes": delivery.stat().st_size,
                },
                ensure_ascii=False,
                indent=2,
            )
        )


if __name__ == "__main__":
    main()
